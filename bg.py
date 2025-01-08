from docx import Document
from docx.shared import Inches
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def add_background_image(doc_path, image_path, output_path):
    """
    Добавляет изображение на задний фон документа Word.
    :param doc_path: Путь к исходному документу Word.
    :param image_path: Путь к изображению.
    :param output_path: Путь для сохранения документа.
    """
    # Пространства имен
    nsmap_v = 'urn:schemas-microsoft-com:vml'
    nsmap_r = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

    # Загружаем документ
    doc = Document(doc_path)

    # Получаем размеры страницы
    section = doc.sections[0]
    page_width = 794 # переводим в пункты (1 inch = 72 points)
    page_height = 1222

    # Получаем заголовок документа
    header = section.header

    # Удаляем старые параграфы в заголовке
    for paragraph in header.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)

    # Создаем временный параграф для добавления изображения
    temp_paragraph = header.add_paragraph()
    run = temp_paragraph.add_run()
    inline_shape = run.add_picture(image_path)
    r_id = inline_shape._inline.graphic.graphicData.pic.blipFill.blip.embed

    # Удаляем временный параграф
    temp_paragraph._element.getparent().remove(temp_paragraph._element)

    # Создаем XML для фонового изображения
    bg_image = parse_xml(
        f"""
        <w:pict {nsdecls('w')} xmlns:v="{nsmap_v}" xmlns:r="{nsmap_r}">
            <v:shape style="position:absolute;left:-76;top:0;width:{page_width};height:{page_height};z-index:-1" 
                coordsize="{page_width},{page_height}">
                <v:imagedata r:id="{r_id}" />
            </v:shape>
        </w:pict>
        """
    )

    # Добавляем XML в заголовок
    header_paragraph = header.add_paragraph()
    header_paragraph._element.append(bg_image)

    # Сохраняем документ
    doc.save(output_path)

# Пример использования
doc_path = "титул_шаблон_пример.docx"  # Исходный шаблон
image_path = "Рисунок1.jpg"  # Путь к изображению
output_path = "bg.docx"  # Новый документ

add_background_image(doc_path, image_path, output_path)

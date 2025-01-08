import os
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from openpyxl import load_workbook
from PIL import Image, ImageTk
from docxtpl import DocxTemplate
from docx2pdf import convert
from reportlab.pdfgen import canvas
from PyPDF2 import PdfReader, PdfWriter
from io import BytesIO
import fitz
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
class ProjectInfoApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Состав проекта")
        self.root.geometry("950x710")

        self.filepath = ""
        self.customer = tk.StringVar()
        self.project_name = tk.StringVar()
        self.project_code = tk.StringVar()  # New field
        self.project_stage = tk.StringVar()
        self.ceo = tk.StringVar()  # New field
        self.project_engineer = tk.StringVar()  # New field
        self.city = tk.StringVar()  # New field
        self.year_of_development = tk.StringVar()  # New field
        self.image_path = tk.StringVar()
        self.dynamic_fields = {}  # Store dynamic fields for the second page

        # Create notebook (tabs)
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill="both", expand=True)

        # Create tabs
        self.tab_file_selection = ttk.Frame(self.notebook)
        self.tab_project_info = ttk.Frame(self.notebook)
        self.tab_titles = ttk.Frame(self.notebook)
        self.tab_iul = ttk.Frame(self.notebook)

        self.notebook.add(self.tab_file_selection, text="Выбор файла")

        # Initially display the file selection tab
        self.create_file_selection_page()
        

    def create_file_selection_page(self):
        # Clear tab content and create content for file selection
        self.clear_tab(self.tab_file_selection)

        btn_select_file = tk.Button(self.tab_file_selection, text="Выберите файл", command=self.select_file)
        btn_select_file.place(relx=0.5, rely=0.5, anchor="center")

    def create_project_info_page(self):
        self.clear_tab(self.tab_project_info)

        # Create left and right frames
        left_frame = tk.Frame(self.tab_project_info)
        left_frame.pack(side="left", fill="both", expand=True, padx=20, pady=20)

        right_frame = tk.Frame(self.tab_project_info)
        right_frame.pack(side="right", fill="both", expand=True, padx=20, pady=20)

        # Left Frame (Fields and Buttons)
        tk.Label(left_frame, text="Заказчик:").pack(anchor="w")
        customer_entry = tk.Entry(left_frame, textvariable=self.customer, width=60)
        customer_entry.pack(anchor="w", pady=5)

        tk.Label(left_frame, text="Наименование проекта:").pack(anchor="w")
        self.project_name_entry = tk.Entry(left_frame, textvariable=self.project_name, width=60)
        self.project_name_entry.pack(anchor="w", pady=5)

        # New fields
        tk.Label(left_frame, text="Шифр проекта:                                                                                                    Этап проекта:").pack(anchor="w")
        code_pack = tk.Frame(left_frame)
        code_pack.pack(anchor="w", pady=5)

        project_code_entry = tk.Entry(code_pack, textvariable=self.project_code, width=53)
        project_code_entry.pack(side= 'left')

        project_stage_entry = tk.Entry(code_pack, textvariable=self.project_stage, width=5)
        project_stage_entry.pack(side= 'left', padx=5)

        tk.Label(left_frame, text="Генеральный директор:").pack(anchor="w")
        ceo_entry = tk.Entry(left_frame, textvariable=self.ceo, width=60)
        ceo_entry.pack(anchor="w", pady=5)

        tk.Label(left_frame, text="Главный инженер проекта:").pack(anchor="w")
        project_engineer_entry = tk.Entry(left_frame, textvariable=self.project_engineer, width=60)
        project_engineer_entry.pack(anchor="w", pady=5)

        tk.Label(left_frame, text="Город:").pack(anchor="w")
        city_entry = tk.Entry(left_frame, textvariable=self.city, width=60)
        city_entry.pack(anchor="w", pady=5)

        tk.Label(left_frame, text="Год разработки:").pack(anchor="w")
        year_of_development_entry = tk.Entry(left_frame, textvariable=self.year_of_development, width=60)
        year_of_development_entry.pack(anchor="w", pady=5)

        btn_select_image = tk.Button(left_frame, text="Выберите изображение для фона", command=self.load_image_path)
        btn_select_image.pack(pady=10, anchor="w")



        btn_save = tk.Button(left_frame, text="Сохранить", command=self.save_to_excel)
        btn_save.pack(pady=10, anchor="w")
        btn_select_image = tk.Button(left_frame, text="Пример титула", command=self.title_example)
        btn_select_image.pack(pady=10, anchor="w")


        # Right Frame (Image Display)
        self.image_panel = tk.Label(right_frame)
        self.image_panel.pack()

        if self.image_path.get():
            self.load_image()

    def create_titles_page(self):
        self.clear_tab(self.tab_titles)

        # Create left and right frames
        left_frame = tk.Frame(self.tab_titles)
        left_frame.pack(side="left", fill="both", expand=True, padx=20, pady=20)

        right_frame = tk.Frame(self.tab_titles)
        right_frame.pack(side="right", fill="both", expand=True, padx=20, pady=20)


        tk.Label(left_frame, text="Выберите из списка:").pack(anchor="w")
        self.dynamic_combobox_titles = ttk.Combobox(left_frame, width=50)
        self.dynamic_combobox_titles.pack(anchor="w", pady=10)
        self.dynamic_combobox_titles.bind("<<ComboboxSelected>>", self.populate_fields_from_selection_titles)

        # Buttons for add, delete, edit, and clear in one row
        button_frame = tk.Frame(left_frame)
        button_frame.pack(pady=10, anchor="w")

        btn_add = tk.Button(button_frame, text="Добавить", command=self.add_to_second_sheet_titles)
        btn_add.pack(side="left", padx=5)

        btn_delete = tk.Button(button_frame, text="Удалить", command=self.delete_selected_item)
        btn_delete.pack(side="left", padx=5)

        btn_edit = tk.Button(button_frame, text="Редактировать", command=self.edit_selected_item)
        btn_edit.pack(side="left", padx=5)

        btn_clear = tk.Button(button_frame, text="Очистить поля", command=self.clear_fields)
        btn_clear.pack(side="left", padx=5)

        # Predefined fields for titles

        self.division_number = tk.StringVar()
        self.division_name = tk.StringVar()
        self.division_code = tk.StringVar()

        self.part_number = tk.StringVar()
        self.part_name = tk.StringVar()

        self.comment = tk.StringVar()

        tk.Label(left_frame, text="Раздел:").pack(anchor="w")
        division_frame = tk.Frame(left_frame)
        division_frame.pack(anchor="w", pady=5)

        division_number_combobox = ttk.Combobox(division_frame, textvariable=self.division_number, width=2)
        division_number_combobox['values'] = [str(i) for i in range(1, 11)]
        division_number_combobox.pack(side="left", padx=5)

        division_name_entry = tk.Entry(division_frame, textvariable=self.division_name, width=46)
        division_name_entry.pack(side="left", padx=5)

        division_code_entry = tk.Entry(division_frame, textvariable=self.division_code, width=5)
        division_code_entry.pack(side="left", padx=3)

        tk.Label(left_frame, text="Часть:").pack(anchor="w")
        part_frame = tk.Frame(left_frame)
        part_frame.pack(anchor="w", pady=5)

        part_number_combobox = ttk.Combobox(part_frame, textvariable=self.part_number, width=2)
        part_number_combobox['values'] = [str(i) for i in range(1, 11)]
        part_number_combobox.pack(side="left", padx=5)

        part_name_entry = tk.Entry(part_frame, textvariable=self.part_name, width=53)
        part_name_entry.pack(side="left", padx=5)

        tk.Label(left_frame, text="Примечание").pack(anchor="w")
        comment_entry = tk.Entry(left_frame, textvariable=self.comment, width=60)
        comment_entry.pack(anchor="w", pady=5)


        btn_create_word_documents = tk.Button(left_frame, text="Создать Word документы", command=self.create_word_documents)
        btn_create_word_documents.pack(anchor="w", pady=10)

        self.update_combobox()


    def create_iul_page(self):
        self.clear_tab(self.tab_iul)

        # Create left and right frames
        left_frame = tk.Frame(self.tab_iul)
        left_frame.pack(side="left", fill="both", expand=True, padx=20, pady=20)

        right_frame = tk.Frame(self.tab_iul)
        right_frame.pack(side="right", fill="both", expand=True, padx=20, pady=20)

        # Create the combobox at the top
        tk.Label(left_frame, text="Выберите из списка:").pack(anchor="w")
        self.dynamic_combobox_iul = ttk.Combobox(left_frame, width=50)
        self.dynamic_combobox_iul.pack(anchor="w", pady=10)
        self.dynamic_combobox_iul.bind("<<ComboboxSelected>>", self.populate_fields_from_selection_iul)

        # Label and button for selecting a PDF file
        self.pdf_path_label = tk.Label(left_frame, text="Путь к pdf файлу:")
        self.pdf_path_label.pack(anchor="w", pady=10)

        btn_select_pdf = tk.Button(left_frame, text="Выберите pdf файл", command=self.load_pdf_file)
        btn_select_pdf.pack(pady=10, anchor="w")


        # Load Excel headers and create fields in right_frame
        try:
            workbook = load_workbook(self.filepath)
            sheet = workbook.worksheets[1]  # Access the second sheet
            headers = [cell.value for cell in sheet[1]]

            # Skip the first 5 headers and start from column 6
            for header in headers[5:]:
                tk.Label(right_frame, text=header).pack(anchor="w")
                entry_var = tk.StringVar()
                entry = tk.Entry(right_frame, textvariable=entry_var, width=60)
                entry.pack(anchor="w", pady=5)

                self.dynamic_fields[header] = entry_var

        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при чтении второго листа: {e}")

        self.update_combobox()

    def add_to_second_sheet_titles(self):
        try:
            workbook = load_workbook(self.filepath)
            sheet = workbook.worksheets[1]

            division_number = self.division_number.get()
            division_name = self.division_name.get()
            division_code = self.division_code.get()
            part_number = self.part_number.get()  # Может быть пустым
            part_name = self.part_name.get()  # Может быть пустым
            comment = self.comment.get()

            if not division_number:
                messagebox.showerror("Ошибка", "Не указан номер раздела")
                return

            # Проверка уникальности комбинации division_number и part_number (если указаны)
            existing_tomes = set(
                (
                    int(sheet.cell(row=i, column=1).value),
                    int(sheet.cell(row=i, column=4).value) if sheet.cell(row=i, column=4).value is not None else None
                )
                for i in range(2, sheet.max_row + 1)  # Пропускаем заголовок (если есть)
                if sheet.cell(row=i, column=1).value is not None
            )

            if (int(division_number), int(part_number) if part_number else None) in existing_tomes:
                messagebox.showerror("Ошибка", "Том с таким номером уже существует")
                return

            # Определение строки для вставки
            insert_row = sheet.max_row + 1
            for i in range(2, sheet.max_row + 1):
                existing_division = sheet.cell(row=i, column=1).value
                existing_part = sheet.cell(row=i, column=4).value

                if existing_division is not None:
                    if int(existing_division) > int(division_number) or (
                        int(existing_division) == int(division_number) and
                        existing_part is not None and
                        int(existing_part) > int(part_number) if part_number else False
                    ):
                        insert_row = i
                        break

            # Сдвиг строк вниз для вставки
            sheet.insert_rows(insert_row)

            # Запись данных в таблицу
            sheet.cell(row=insert_row, column=1).value = int(division_number)
            sheet.cell(row=insert_row, column=2).value = division_name
            sheet.cell(row=insert_row, column=3).value = division_code
            sheet.cell(row=insert_row, column=4).value = int(part_number) if part_number else None
            sheet.cell(row=insert_row, column=5).value = part_name
            sheet.cell(row=insert_row, column=6).value = comment

            workbook.save(self.filepath)
            messagebox.showinfo("Успех", "Данные успешно сохранены в Excel файл.")

            self.update_combobox()

            new_value = f"Том {division_number}"
            if part_number:
                new_value += f".{part_number} - {part_name}"
            else:
                new_value += f" - {division_name}"

            # Установка нового значения в комбобоксе
            self.dynamic_combobox_titles.set(new_value)

        except ValueError as ve:
            messagebox.showerror("Ошибка", f"Некорректные данные: {ve}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при записи в файл: {e}")



    def delete_selected_item(self):
            selected_value = self.dynamic_combobox_titles.get()

            if selected_value:
                try:
                    if '.' in selected_value.split('Том ')[1].split(' - ')[0]:
                        division_number = selected_value.split('Том ')[1].split('.')[0]
                        part_number = selected_value.split('.')[1].split(' - ')[0]
                    else:
                        division_number = selected_value.split('Том ')[1].split(' - ')[0]
                        part_number = None                       

                    workbook = load_workbook(self.filepath)
                    sheet = workbook.worksheets[1]

                    for row in range(2, sheet.max_row + 1):
                        existing_division = str(sheet.cell(row=row, column=1).value)
                        existing_part = str(sheet.cell(row=row, column=4).value) if sheet.cell(row=row, column=4).value is not None else None

                        if existing_division == division_number and existing_part == part_number:
                            sheet.delete_rows(row)
                            break
                    else:
                        messagebox.showerror("Ошибка", f"Том {division_number}{f'.{part_number}' if part_number else ''} не найден")

                    workbook.save(self.filepath)
                    messagebox.showinfo("Успех", "Выбранная строка успешно удалена")

                    self.update_combobox()
                    self.clear_fields()
                except Exception as e:
                    messagebox.showerror("Ошибка", f"Ошибка при удалении строки: {e}")

    def edit_selected_item(self):
        try:
            workbook = load_workbook(self.filepath)
            sheet = workbook.worksheets[1]

            division_number = self.division_number.get()
            part_number = self.part_number.get()  # Может быть пустым

            if not division_number:
                messagebox.showerror("Ошибка", "Не указан номер раздела.")
                return

            division_number = int(division_number)
            part_number = int(part_number) if part_number else None

            # Ищем строку для редактирования
            row_to_edit = None
            for i in range(2, sheet.max_row + 1):
                existing_division = sheet.cell(row=i, column=1).value
                existing_part = sheet.cell(row=i, column=4).value

                if existing_division is not None:
                    if int(existing_division) == division_number and (int(existing_part) == part_number if part_number else existing_part is None):
                        row_to_edit = i
                        break

            if row_to_edit is None:
                messagebox.showerror("Ошибка", "Том с указанными номерами раздела и части не найден.")
                return

            # Обновляем данные в найденной строке
            sheet.cell(row=row_to_edit, column=2).value = self.division_name.get()
            sheet.cell(row=row_to_edit, column=3).value = self.division_code.get()
            sheet.cell(row=row_to_edit, column=5).value = self.part_name.get()
            sheet.cell(row=row_to_edit, column=6).value = self.comment.get()

            # Сохраняем файл
            workbook.save(self.filepath)
            messagebox.showinfo("Успех", "Данные успешно обновлены в Excel файле.")

            self.update_combobox()


            new_value = f"Том {division_number}"
            if part_number:
                new_value += f".{part_number} - {self.part_name.get()}"
            else:
                new_value += f" - {self.division_name.get()}"

            # Установка нового значения в комбобоксе
            self.dynamic_combobox_titles.set(new_value)
        except ValueError as ve:
            messagebox.showerror("Ошибка", f"Некорректные данные: {ve}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при редактировании файла: {e}")


    def clear_fields(self):
        self.division_number.set("")
        self.division_name.set("")
        self.division_code.set("")

        self.part_number.set("")
        self.part_name.set("")

        self.comment.set("")
        self.dynamic_combobox_titles.set("")

    def clear_tab(self, tab):
        for widget in tab.winfo_children():
            widget.destroy()
    def load_image_path(self):
        file_path = filedialog.askopenfilename(filetypes=[("Image files", "*.jpg"), ("Image files", "*.png")])
        if file_path:
            self.image_path.set(file_path)
            self.load_image()
    def load_pdf_file(self):
        selected_value = self.dynamic_combobox_iul.get()
        if not selected_value:
            messagebox.showwarning("Предупреждение", "Пожалуйста, выберите элемент из списка.")
            return

        file_path = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf")])
        if file_path:
            tome_number = selected_value.split(" - ")[0]
            try:
                workbook = load_workbook(self.filepath)
                sheet = workbook.worksheets[1]

                for row in range(2, sheet.max_row + 1):
                    if str(sheet.cell(row=row, column=1).value) == tome_number:
                        sheet.cell(row=row, column=5).value = file_path
                        break

                workbook.save(self.filepath)
                messagebox.showinfo("Успех", "Путь к PDF файлу успешно сохранён в Excel файл.")

                # Update the PDF path label with the new file path
                self.pdf_path_label.config(text=f"Путь к pdf файлу: {file_path}")
            except Exception as e:
                messagebox.showerror("Ошибка", f"Ошибка при записи файла: {e}")

    def select_file(self):
        self.filepath = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
        if self.filepath:
            self.load_excel_data()
            self.create_project_info_page()
            self.create_titles_page()
            self.create_iul_page()
            self.notebook.forget(self.tab_file_selection)
            self.notebook.add(self.tab_project_info, text="Информация о проекте")
            self.notebook.add(self.tab_titles, text="Состав проекта")
            self.notebook.add(self.tab_iul, text="Создание ИУЛ")

            
        try:
            workbook = load_workbook(self.filepath)
            if "Сведения о проекте" not in workbook.sheetnames:
                sheet = workbook.create_sheet("Сведения о проекте")
                sheet.append(["Заказчик", "Наименование проекта", "Фон", "Шифр проекта","Стадия проектирования", "Генеральный директор", 
                              "Главный инженер проекта", "Город", "Год разработки"])
            if "Состав проекта" not in workbook.sheetnames:
                sheet = workbook.create_sheet("Состав проекта")
                sheet.append(["Номер раздела", "Наименование раздела", "Обозначение раздела", "Номер части", "Наименование части", "Примечание", "pdf-файл", "Разработал", "Проверил", "ГИП", "Н.Контроль"])
            if workbook.sheetnames[0] != "Сведения о проекте" and workbook.sheetnames[0] != "Состав проекта":
                del workbook[workbook.sheetnames[0]]

                
                workbook.save(self.filepath)
        except Exception as e:
                messagebox.showerror("Ошибка", f"Ошибка при создании листов: {e}")

    def load_excel_data(self):
        try:
            workbook = load_workbook(self.filepath)
            sheet = workbook.worksheets[0]  # Access the first sheet
            self.customer.set(sheet["A2"].value or "")
            self.project_name.set(sheet["B2"].value or "")

            self.image_path.set(sheet["C2"].value or "")
            self.project_code.set(sheet["D2"].value or "")
            self.project_stage.set(sheet["E2"].value or "")
            self.ceo.set(sheet["F2"].value or "")
            self.project_engineer.set(sheet["G2"].value or "")
            self.city.set(sheet["H2"].value or "")
            self.year_of_development.set(sheet["I2"].value or "")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при чтении файла: {e}")
            print(e)

    def load_image(self):
        try:
            if self.image_path.get():
                img = Image.open(self.image_path.get())
                img = img.resize((465, 657), Image.LANCZOS)
                img = ImageTk.PhotoImage(img)
                self.image_panel.config(image=img)
                self.image_panel.image = img
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при загрузке изображения: {e}")
    def title_example(self):
        output_path_docx = "temptitle.docx"
        output_path_pdf = "temptitle.pdf"
        output_path_png = "temptitle.png"

        try:
            doc = DocxTemplate("титул_шаблон.docx")

            replacements = {
                'Заказчик': self.customer.get(),
                'Наименование': self.project_name.get(),
                'Шифр': f"{self.project_code.get()}-{self.project_stage.get()}",
                'Генеральный_директор': self.ceo.get(),
                'Главный_инженер_проекта': self.project_engineer.get(),
                'Город': self.city.get(),
                'Год': self.year_of_development.get(),
                'Номер_раздела': 0,
                'Наименование_раздела': 'Проектная документация',
                'Часть': 'Часть 0 «Пример»',
                'Том': 'Том 0.0'
            }

            doc.render(replacements)

            doc.save(output_path_docx)

            
            if self.image_path.get():
                self.add_background_image(output_path_docx, output_path_docx)

            convert(output_path_docx, output_path_pdf)



            pdf_document = fitz.open(output_path_pdf)
            page = pdf_document.load_page(0)
            pix = page.get_pixmap()
            pix.save(output_path_png)

            img = Image.open(output_path_png)
            img = img.resize((465, 657), Image.LANCZOS)
            img = ImageTk.PhotoImage(img)
            self.image_panel.config(image=img)
            self.image_panel.image = img

        finally:
            for temp_file in [output_path_docx, output_path_pdf,output_path_png]:
                try:
                    if os.path.exists(temp_file):
                        os.remove(temp_file)
                except Exception as e:
                    print(f"Не удалось удалить {temp_file}: {e}")

    def add_background_image(self, doc_path, output_path):
        # Пространства имен
        nsmap_v = 'urn:schemas-microsoft-com:vml'
        nsmap_r = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

        # Загружаем документ
        doc = Document(doc_path)

        # Получаем размеры страницы
        section = doc.sections[0]
        page_width = 794  # переводим в пункты (1 inch = 72 points)
        page_height = 1122

        # Получаем заголовок документа
        header = section.header

        # Удаляем старые параграфы в заголовке
        for paragraph in header.paragraphs:
            p = paragraph._element
            p.getparent().remove(p)

        # Создаем временный параграф для добавления изображения
        temp_paragraph = header.add_paragraph()
        run = temp_paragraph.add_run()
        inline_shape = run.add_picture(self.image_path.get())
        r_id = inline_shape._inline.graphic.graphicData.pic.blipFill.blip.embed

        # Удаляем временный параграф
        temp_paragraph._element.getparent().remove(temp_paragraph._element)

        # Создаем XML для фонового изображения
        bg_image = parse_xml(
            f"""
            <w:pict {nsdecls('w')} xmlns:v="{nsmap_v}" xmlns:r="{nsmap_r}">
                <v:shape style="position:absolute;left:-95;top:0;width:{page_width};height:{page_height};z-index:-1" 
                    coordsize="{page_width},{page_height}">
                    <v:imagedata r:id="{r_id}" />
                </v:shape>
            </w:pict>
            """
        )

        # Добавляем XML в заголовок
        header_paragraph = header.add_paragraph()
        header_paragraph._element.append(bg_image)

        # Сохраняем документ
        doc.save(output_path)

    def populate_fields_from_selection_titles(self, event):

        selected_value = self.dynamic_combobox_titles.get()

        if selected_value:
            if '.' in selected_value.split('Том ')[1].split(' - ')[0]:
                division_number = selected_value.split('Том ')[1].split('.')[0]
                part_number = selected_value.split('.')[1].split(' - ')[0]
            else:
                division_number = selected_value.split('Том ')[1].split(' - ')[0]
                part_number = None
            try:
                workbook = load_workbook(self.filepath)
                sheet = workbook.worksheets[1]

                # Iterate over rows to find the selected "Номер тома"
                for row in sheet.iter_rows(min_row=2, values_only=True):
                    if str(row[0]) == str(division_number) and str(row[3]) == str(part_number):
                        # Populate only the first 4 dynamic fields based on the row data
                        self.division_number.set(row[0] if row[0] else "") 
                        self.division_name.set(row[1] if row[1] else "")
                        self.division_code.set(row[2] if row[2] else "")

                        self.part_number.set(row[3] if row[3] else "")  
                        self.part_name.set(row[4] if row[4] else "")
                        self.comment.set(row[5] if row[5] else "")
            except Exception as e:
                messagebox.showerror("Ошибка", f"Ошибка при заполнении полей: {e}")
    def populate_fields_from_selection_iul(self, event):

        selected_value = self.dynamic_combobox_iul.get()

        if selected_value:
            tome_number = selected_value.split(" - ")[0]
            try:
                workbook = load_workbook(self.filepath)
                sheet = workbook.worksheets[1]

                # Iterate over rows to find the selected "Номер тома"
                for row in sheet.iter_rows(min_row=2, values_only=True):
                    if str(row[0]) == tome_number:
                        # Populate only the first 4 dynamic fields based on the row data
                        for i, (header, entry_var) in enumerate(self.dynamic_fields.items()):
                            if i >= 4:
                                entry_var.set(row[i+1] if i <= len(row) else "")
                        break
            except Exception as e:
                messagebox.showerror("Ошибка", f"Ошибка при заполнении полей: {e}")
    def save_to_excel(self):
        try:
            workbook = load_workbook(self.filepath)
            sheet = workbook.worksheets[0]  # Access the first sheet
            sheet["A2"] = self.customer.get()
            sheet["B2"] = self.project_name.get()
            sheet["C2"] = self.image_path.get()
            sheet["D2"] = self.project_code.get()  # Save new field
            sheet["E2"] = self.project_stage.get()  # Save new field
            sheet["F2"] = self.ceo.get()  # Save new field
            sheet["G2"] = self.project_engineer.get()  # Save new field
            sheet["H2"] = self.city.get()  # Save new field
            sheet["I2"] = self.year_of_development.get()  # Save new field

            workbook.save(self.filepath)
            messagebox.showinfo("Успех", "Данные успешно сохранены в Excel файл.")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при сохранении файла: {e}")

    def update_combobox(self):
        try:
            workbook = load_workbook(self.filepath)
            sheet = workbook.worksheets[1]  # Access the second sheet
            rows = []
            for row in sheet.iter_rows(min_row=2, values_only=True):
                if row[0]:  # Ensure there's data in the division number column
                    tom_value = f"Том {row[0]}"
                    if row[3]:  # Check if part number exists
                        tom_value += f".{row[3]} - {row[4]}"
                    else:  # Check if part name exists
                        tom_value += f" - {row[1]}"
                    rows.append(tom_value)

            if hasattr(self, 'dynamic_combobox_iul'):
                self.dynamic_combobox_iul["values"] = rows
            self.dynamic_combobox_titles["values"] = rows
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при обновлении списка: {e}")
    def create_word_documents(self):
        try:
            workbook = load_workbook(self.filepath)
            sheet = workbook.worksheets[1]  # Access the second sheet
            doc = DocxTemplate("титул_шаблон.docx")
            for row in sheet.iter_rows(min_row=2, values_only=True):
                if row[0]:  
                    folder_path = f"Проектная документация/Раздел ПД №{row[0]}"
                    part = ''
                    tome = f"Том {row[0]}"
                    if row[3]:  
                        folder_path += f" подраздел ПД №{row[0]}.{row[3]}"
                        part = f"Часть {row[3]} «{row[4]}»"
                        tome += f".{row[3]}"

                    replacements = {
                        'Заказчик': self.customer.get(),
                        'Наименование': self.project_name.get(),
                        'Шифр': f"{self.project_code.get()}-{self.project_stage.get()}-{row[2]}{row[3]if row[3] else ''}",
                        'Генеральный_директор': self.ceo.get(),
                        'Главный_инженер_проекта': self.project_engineer.get(),
                        'Город': self.city.get(),
                        'Год': self.year_of_development.get(),
                        'Номер_раздела': row[0],
                        'Наименование_раздела': row[1],
                        'Часть': part,
                        'Том': tome
                        
                    }
                    os.makedirs(folder_path, exist_ok=True)
                    filename = f"{folder_path}/00_Титул_{tome}.docx"
                    doc.render(replacements)
                    doc.save(filename)

                    if self.image_path.get():
                        self.add_background_image(filename, filename)
            messagebox.showinfo("Успех", f"Документы созданы в папке: Проектная документация")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при создании документов: {e}")
root = tk.Tk()
app = ProjectInfoApp(root)
root.mainloop()
